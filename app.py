import gradio as gr
import os
import docx
import numpy as np
from sentence_transformers import SentenceTransformer
from sklearn.metrics.pairwise import cosine_similarity
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain_community.llms import HuggingFaceEndpoint
from langchain_huggingface import HuggingFaceEmbeddings
import torch
from PIL import Image
from torchvision import transforms
from torchvision.models import resnet50, ResNet50_Weights
from torchvision import transforms, models
from sentence_transformers import CrossEncoder


reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-6-v2')

class GeometryImageClassifier:
    def __init__(self):
        # Load ResNet50 but only use it for feature extraction
        self.model = models.resnet50(weights='DEFAULT')
        # Remove the final classification layer
        self.model.fc = torch.nn.Identity()
        self.model.eval()
        
        self.transform = transforms.Compose([
            transforms.Resize((224, 224)),
            transforms.ToTensor(),
            transforms.Normalize(mean=[0.485, 0.456, 0.406], std=[0.229, 0.224, 0.225])
        ])
        
        self.reference_embeddings = {
            "flat.png": {
                "embedding": None,
                "label": "Flat or Sheet-Based"
            },
            "flat_angle2.png": {
                "embedding": None,
                "label": "Flat or Sheet-Based"
            },
            "cylindrical.png": {
                "embedding": None,
                "label": "Cylindrical"
            },
            "cylindrical_angle2.png": {
                "embedding": None,
                "label": "Cylindrical"
            },
            "complex.png": {
                "embedding": None,
                "label": "Complex Multi Axis Geometry"
            },
            "complex_angle2.png": {
                "embedding": None,
                "label": "Complex Multi Axis Geometry"
            }
        }
        
    def compute_embedding(self, images):
        img = Image.open(images).convert('RGB')
        img_tensor = self.transform(img).unsqueeze(0)
        
        with torch.no_grad():
            embedding = self.model(img_tensor)
        return embedding.squeeze().numpy()
    
    def initialize_reference_embeddings(self, reference_folder):
        for image_name in self.reference_embeddings.keys():
            images = f"{reference_folder}/{image_name}"
            self.reference_embeddings[image_name]["embedding"] = self.compute_embedding(images)
    
    def find_closest_geometry(self, query_embedding):
        best_similarity = -1
        best_label = None
        
        for ref_data in self.reference_embeddings.values():
            similarity = cosine_similarity(
                query_embedding.reshape(1, -1),
                ref_data["embedding"].reshape(1, -1)
            )[0][0]
            
            if similarity > best_similarity:
                best_similarity = similarity
                best_label = ref_data["label"]
        
        return best_label
    
    def process_image(self, images):
        # Compute embedding for the input image
        query_embedding = self.compute_embedding(images)
        
        # Find the closest matching reference geometry
        return self.find_closest_geometry(query_embedding)


    
# ✅ Use a strong sentence embedding model
semantic_model = SentenceTransformer("all-MiniLM-L6-v2")

def extract_text_from_docx(file_path):
    """ ✅ Extracts normal text & tables from a .docx file for better retrieval. """
    doc = docx.Document(file_path)
    extracted_text = []


    for para in doc.paragraphs:
        if para.text.strip():
            extracted_text.append(para.text.strip())


    for table in doc.tables:
        extracted_text.append("📌 Table Detected:")
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            if any(row_text):
                extracted_text.append(" | ".join(row_text))


    return "\n".join(extracted_text)

def load_documents():
    """ ✅ Loads & processes documents, ensuring table data is properly extracted. """
    file_paths = {
        "Fastener_Types_Manual": "Fastener_Types_Manual.docx",
        "Manufacturing_Expert_Manual": "Manufacturing Expert Manual.docx"
    }
    all_splits = []

    for doc_name, file_path in file_paths.items():
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Document not found: {file_path}")


        print(f"Extracting text from {file_path}...")
        full_text = extract_text_from_docx(file_path)


        text_splitter = RecursiveCharacterTextSplitter(chunk_size=1500, chunk_overlap=200)
        doc_splits = text_splitter.create_documents([full_text])


        for chunk in doc_splits:
            chunk.metadata = {"source": doc_name}


        all_splits.extend(doc_splits)


    return all_splits

def create_db(splits):
    """ ✅ Creates a FAISS vector database from document splits. """
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-base-en-v1.5")
    vectordb = FAISS.from_documents(splits, embeddings)
    return vectordb

def rerank_documents(query, docs, top_k=5):
    pairs = [[query, doc.page_content] for doc in docs]
    scores = reranker.predict(pairs)
    doc_score_pairs = list(zip(docs, scores))
    ranked_docs = sorted(doc_score_pairs, key=lambda x: x[1], reverse=True)
    return [doc for doc, score in ranked_docs[:top_k]]

def filter_relevant_chunks(query, chunks, embeddings, threshold=0.5):
    query_embedding = embeddings.embed_query(query)
    filtered_chunks = []
    for chunk in chunks:
        chunk_embedding = embeddings.embed_query(chunk.page_content)
        similarity = cosine_similarity([query_embedding], [chunk_embedding])[0][0]
        if similarity > threshold:
            filtered_chunks.append(chunk)
    return filtered_chunks

def retrieve_documents(query, retriever, embeddings):
    print("\n=== Document Retrieval Process ===")
    print(f"Query: {query}")
    
    results = retriever.invoke(query)
    print(f"Initial results count: {len(results)}")
    
    if not results:
        print("No initial results found")
        return []
    
    reranked_results = rerank_documents(query, results, top_k=5)
    print(f"Reranked results count: {len(reranked_results)}")
    
    filtered_chunks = filter_relevant_chunks(query, reranked_results, embeddings, threshold=0.3)
    print(f"Filtered chunks count: {len(filtered_chunks)}")
    
    if not filtered_chunks:
        print("No chunks passed filtering")
        return []
    
    doc_embeddings = np.array([embeddings.embed_query(doc.page_content) for doc in filtered_chunks])
    query_embedding = np.array(embeddings.embed_query(query)).reshape(1, -1)
    similarity_scores = cosine_similarity(query_embedding, doc_embeddings)[0]
    
    print("\nSimilarity Scores:")
    for doc, score in zip(filtered_chunks, similarity_scores):
        print(f"Score: {score:.4f} | Source: {doc.metadata.get('source', 'Unknown')}")
        print(f"Content Preview: {doc.page_content[:100]}...\n")
    
    MIN_SIMILARITY = 0.3
    filtered_results = [(doc, sim) for doc, sim in zip(filtered_chunks, similarity_scores) if sim >= MIN_SIMILARITY]
    print(f"Final filtered results count: {len(filtered_results)}")
    
    return [doc for doc, _ in filtered_results] if filtered_results else []

def validate_query_semantically(query, retrieved_docs):
    print("\n=== Semantic Validation ===")
    if not retrieved_docs:
        print("No documents to validate")
        return False

    combined_text = " ".join([doc.page_content for doc in retrieved_docs])
    query_embedding = semantic_model.encode(query, normalize_embeddings=True)
    doc_embedding = semantic_model.encode(combined_text, normalize_embeddings=True)
    similarity_score = np.dot(query_embedding, doc_embedding)
    
    print(f"Query: {query}")
    print(f"Semantic similarity score: {similarity_score:.4f}")
    print(f"Validation {'passed' if similarity_score >= 0.3 else 'failed'}")
    
    return similarity_score >= 0.3

def handle_query(query, history, retriever, qa_chain, embeddings):
    """ ✅ Handles user queries & prevents hallucination. """
    retrieved_docs = retrieve_documents(query, retriever, embeddings)
    if not retrieved_docs or not validate_query_semantically(query, retrieved_docs):
        return history + [(query, "I couldn't find any relevant information.")], ""
    response = qa_chain.invoke({"question": query, "chat_history": history})
    assistant_response = response['answer'].strip()
    if not validate_query_semantically(query, retrieved_docs):
        assistant_response = "I couldn't find any relevant information."
    assistant_response += f"\n\n📄 **Source:** {', '.join(set(doc.metadata.get('source', 'Unknown') for doc in retrieved_docs))}"
    print(f"🤖 LLM Response: {assistant_response[:300]}")  # ✅ Limit output for debugging
    history.append((query, assistant_response))
    return history, ""

def initialize_chatbot(vector_db):
    """ ✅ Initializes chatbot with improved retrieval & processing. """
    memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True, output_key='answer')
    embeddings = HuggingFaceEmbeddings(model_name="BAAI/bge-base-en-v1.5")
    retriever = vector_db.as_retriever(search_kwargs={"k": 5, "search_type": "similarity"})
    system_prompt = """You are an AI assistant that answers questions **ONLY based on the provided documents**.
- **If no relevant documents are retrieved, respond with: "I couldn't find any relevant information."**
- **If the meaning of the query does not match the retrieved documents, say "I couldn't find any relevant information."**
- **Do NOT attempt to answer from general knowledge.**
"""
    llm = HuggingFaceEndpoint(
        repo_id="tiiuae/falcon-40b-instruct",
        huggingfacehub_api_token=os.environ.get("HUGGINGFACE_API_TOKEN"),
        temperature=0.1,
        max_new_tokens=400,  
        task="text-generation",
        system_prompt=system_prompt)

    qa_chain = ConversationalRetrievalChain.from_llm(
        llm=llm,
        retriever=retriever,
        memory=memory,
        return_source_documents=True,
        verbose=False)
    return retriever, qa_chain, embeddings


def process_image_and_generate_query(image):
    classifier = GeometryImageClassifier()
    geometry_type = classifier.process_image(image)
    
    query = f"I have a {geometry_type} geometry, which screw should I use and what is the best machine to use for {geometry_type} geometry?"
    return geometry_type, query


def demo():
    # Initialize classifier once at startup
    classifier = GeometryImageClassifier()
    classifier.initialize_reference_embeddings("images")
    
    # Initialize chatbot components
    retriever, qa_chain, embeddings = initialize_chatbot(create_db(load_documents()))
    
    with gr.Blocks() as app:
        gr.Markdown("### 🤖 **Fastener Agent with Image Recognition** 📚")
        
        with gr.Row():
            with gr.Column(scale=1):
                image_input = gr.Image(type="filepath", label="Upload Geometry Image")
                geometry_label = gr.Textbox(label="Detected Geometry Type", interactive=False)
                
            with gr.Column(scale=2):
                chatbot = gr.Chatbot()
                query_input = gr.Textbox(label="Ask me a question")
                query_btn = gr.Button("Ask")


        def image_upload_handler(image):
            if image is None:
                return "", ""
            # Use the initialized classifier
            geometry_type = classifier.process_image(image)
            suggested_query = f"I have a {geometry_type} geometry, which screw should I use and what is the best machine to use for {geometry_type} geometry?"
            return geometry_type, suggested_query


        def user_query_handler(query, history):
            return handle_query(query, history, retriever, qa_chain, embeddings)


        image_input.change(
            image_upload_handler,
            inputs=[image_input],
            outputs=[geometry_label, query_input]
        )
        
        query_btn.click(
            user_query_handler,
            inputs=[query_input, chatbot],
            outputs=[chatbot, query_input]
        )
        
        query_input.submit(
            user_query_handler,
            inputs=[query_input, chatbot],
            outputs=[chatbot, query_input]
        )


    app.launch()


if __name__ == "__main__":
    demo()
